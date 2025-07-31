import os
import re
from flask import Flask, render_template, request, send_from_directory
from docx import Document
from docx.shared import Inches
from datetime import datetime

app = Flask(__name__)

# 配置路径
PERSON_DOCS_DIR = 'person_docs'
OUTPUT_DIR = 'output'
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# 有效期检查函数
def extract_expiry_from_filename(filename):
    match = re.search(r'(\d{4}-\d{2}-\d{2})', filename)
    if match:
        try:
            return datetime.strptime(match.group(1), '%Y-%m-%d').date()
        except:
            return None
    return None

def is_expired(expiry_date):
    if expiry_date:
        return expiry_date < datetime.today().date()
    return False

# 插入图片到 Word（每张单独一页）
def insert_person_images(doc, name, role):
    folder_path = os.path.join(PERSON_DOCS_DIR, name)
    if not os.path.exists(folder_path):
        doc.add_paragraph(f"⚠️ 未找到 {name} 的资料文件夹")
        return

    # 所有图片按顺序插入
    image_files = sorted(os.listdir(folder_path))
    
    # 插入身份证
    for img in image_files:
        if "身份证" in img:
            expiry = extract_expiry_from_filename(img)
            p = doc.add_paragraph(f"{role} {name} - 身份证")
            if expiry and is_expired(expiry):
                p.add_run(" ⚠️ 有效期已过")
            doc.add_picture(os.path.join(folder_path, img), width=Inches(6))
            doc.add_page_break()
            break

    # 插入中/高工证
    for level in ["中级工程师", "高级工程师"]:
        for img in image_files:
            if level in img:
                doc.add_paragraph(f"{name} - {level}")
                doc.add_picture(os.path.join(folder_path, img), width=Inches(6))
                doc.add_page_break()

    # 插入岗位对应证书
    required_keywords = {
        "项目经理": ["建造师", "B证"],
        "技术负责人": [],
        "质量员": ["质量员"],
        "安全员": ["C证"]
    }
    for keyword in required_keywords.get(role, []):
        for img in image_files:
            if keyword in img:
                expiry = extract_expiry_from_filename(img)
                p = doc.add_paragraph(f"{role} {name} - {keyword}")
                if expiry and is_expired(expiry):
                    p.add_run(" ⚠️ 有效期已过")
                doc.add_picture(os.path.join(folder_path, img), width=Inches(6))
                doc.add_page_break()

    # 插入劳动合同与社保
    for tag in ["劳动合同", "社保"]:
        for img in image_files:
            if tag in img:
                doc.add_paragraph(f"{name} - {tag}")
                doc.add_picture(os.path.join(folder_path, img), width=Inches(6))
                doc.add_page_break()

@app.route('/', methods=['GET', 'POST'])
def index():
    file_link = None
    if request.method == 'POST':
        pm = request.form['pm']
        tech = request.form['tech']
        quality = request.form['quality']
        safety = request.form['safety']

        # 创建Word文档
        doc = Document()
        doc.add_heading("班子人员资料（自动生成）", 0)

        insert_person_images(doc, pm, "项目经理")
        insert_person_images(doc, tech, "技术负责人")
        insert_person_images(doc, quality, "质量员")
        insert_person_images(doc, safety, "安全员")

        filename = f"标书人员附件_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(OUTPUT_DIR, filename)
        doc.save(filepath)

        file_link = f"/download/{filename}"

    return render_template('index.html', file_link=file_link)

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
